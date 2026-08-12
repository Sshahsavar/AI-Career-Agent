import streamlit as st
import json
import os
from datetime import datetime
import pandas as pd
import requests
from bs4 import BeautifulSoup
from google import genai
from google.genai import types
from google.genai.errors import ClientError
from pydantic import BaseModel, Field
from typing import List
from fpdf import FPDF
import re
from dotenv import load_dotenv

# Import python-docx for Word document generation
try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("Please install python-docx: pip install python-docx")

# Load local environment variables from .env file
load_dotenv()

# --- INITIALIZATION: AUTO-CREATE FILES FOR GITHUB COMPATIBILITY ---
def initialize_system():
    os.makedirs("profile_data", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    os.makedirs("templates", exist_ok=True)

    # 1. Initialize GitHub Projects JSON
    if not os.path.exists("profile_data/github_repos.json"):
        with open("profile_data/github_repos.json", "w") as f:
            json.dump([], f)
            
    # 2. Initialize Master Resume JSON (PII SCRUBBED)
    if not os.path.exists("profile_data/resume_data.json"):
        default_resume = {
            "header": "# [YOUR NAME]\n[CITY, STATE] | [PHONE] | [EMAIL] | [LINKEDIN] | [GITHUB]",
            "summary": "Quantitative Financial Analyst...",
            "education": "**University Name | Degree (2025–2026)**\nSelected Coursework: ...",
            "experience": "**Company Name | Job Title (2017–2023)**\n* Coordinated research projects...\n* Managed survey data...",
            "skills": "* **Languages & Databases:** Python, R, SQL\n* **Machine Learning:** Random Forest, LASSO"
        }
        with open("profile_data/resume_data.json", "w") as f:
            json.dump(default_resume, f, indent=4)

    # 3. Initialize Soft Skills JSON
    if not os.path.exists("profile_data/soft_skills.json"):
        default_soft_skills = [
            {
                "trait": "Adaptability Under Ambiguity",
                "example": "Navigated severe operational and financial disruptions as an operations manager during the COVID 19 pandemic, rapidly adjusting workflows, inventory controls, and team priorities under intense pressure."
            },
            {
                "trait": "Cross Disciplinary Collaboration",
                "example": "Spearheaded public policy conferences and coordinated multi disciplinary social science research projects, aligning diverse stakeholders across academic, administrative, and policy domains."
            },
            {
                "trait": "Analytical Rigor & Critical Skepticism",
                "example": "Developed deep quantitative intuition through a Masters in Economics at York University, evaluating complex time series dynamics, Copula EGARCH models, and GARCH frameworks to isolate market regime changes."
            },
            {
                "trait": "Strategic Agility",
                "example": "Managed retail operations and cash flow at 5th Garden Coffee Shop during economic downturns, proactively recalibrating pricing models and cost structures to cushion inflationary impacts."
            },
            {
                "trait": "Narrative Synthesis & Communication",
                "example": "Produced and edited 20 full podcast episodes covering complex topics in human development and economic inequality, translating intricate research into accessible, engaging audio content for a broad audience."
            }
        ]
        with open("profile_data/soft_skills.json", "w") as f:
            json.dump(default_soft_skills, f, indent=4)

    # 4. Initialize Dynamic AI Prompts JSON
    if not os.path.exists("profile_data/prompts.json"):
        default_prompts = {
            "main_prompt": """You are an expert career agent. Execute these steps strictly based on the provided Job Description, Master Resume, GitHub Repositories, and Soft Skills profile.

1. MINE KEYWORDS: Extract core skills and requirements from the Job Description.
{project_instruction}
3. OPTIMIZE RESUME: Optimize ONLY the 'Professional Summary' section to highlight relevant skills. DO NOT change Experience or Education.
4. MATCH & GAP ANALYSIS: Determine success probability, list clear matches, and list specific gaps.
5. GENERATE COVER LETTER: Using the attached CV, Soft Skills profile, and JD, write a cover letter following this exact four paragraph structure:

Paragraph 1: Open immediately by identifying the most important or unusual keyword in the job description. Define it as the company's primary objective, explain why they focus on it, and state your desire to contribute to that specific objective. (DO NOT use "I am writing to express..." or any standard application boilerplate).
Paragraph 2: Draw directly from the user's Soft Skills profile below, identifying the psychological mindset or cognitive trait and weaving it together with their exact concrete examples.
Paragraph 3: Address just TWO primary technical challenges or abilities from the job description by integrating specific tools, projects, and technical details found exclusively in the CV, without using phrases like "The primary technical challenge is...".
Paragraph 4: Conclude by expressing your interest in discussing the role because you genuinely appreciate their team, without mentioning any city, location, or physical environment.

Style and Constraints for Cover Letter:
- Write in a completely natural, conversational, human voice. Vary your sentence lengths and use active voice.
- ABSOLUTE BAN ON AI BOILERPLATE: Do NOT use phrases like "I am writing to express", "Furthermore", "In addition", "For instance", "As a [Role]", or "The primary technical challenge".
- STRICT NEGATIVE CONSTRAINT: Do NOT use hyphens or dashes (-) at all anywhere in the text. Use unhyphenated alternatives.
- Fact check everything against the attached CV and Soft Skills profile. Never invent or assume.""",
            
            "regen_prompt": """Instructions for Cover Letter Regeneration:
- Maintain the exact four paragraph structure:
  Paragraph 1: Open immediately with keyword, define primary objective, explain why, state desire to contribute (NO boilerplate).
  Paragraph 2: Draw directly from the Soft Skills & Cognitive Abilities profile, weaving the psychological mindset with the user's exact concrete examples.
  Paragraph 3: TWO primary technical challenges or abilities mapped to tools/projects in CV (NO phrases like "The primary technical challenge is...").
  Paragraph 4: Conclude expressing interest because of team/appreciation, NO city/location/physical environment.
- Completely natural, conversational, human voice. Vary sentence lengths, use active voice.
- ABSOLUTE BAN ON AI BOILERPLATE: Do NOT use phrases like "I am writing to express", "Furthermore", "In addition", "For instance", "As a [Role]", or "The primary technical challenge".
- STRICT NEGATIVE CONSTRAINT: Do NOT use hyphens or dashes (-) at all anywhere in the text. Use unhyphenated alternatives.
- Fact check everything against the attached CV and Soft Skills. Never invent, assume, or include any experience, skill, or metric that is not explicitly written.
- Return ONLY the revised cover letter body text as a plain string."""
        }
        with open("profile_data/prompts.json", "w") as f:
            json.dump(default_prompts, f, indent=4)

initialize_system()

# --- TEXT SANITIZATION & MINIFICATION HELPERS ---
def sanitize_text(text):
    if not isinstance(text, str):
        return text
    return (
        text.replace("–", "-")
            .replace("—", "-")
            .replace("−", "-")
            .replace("“", '"')
            .replace("”", '"')
            .replace("‘", "'")
            .replace("’", "'")
    )

def minify_text_for_llm(text):
    return re.sub(r'\s+', ' ', text).strip()

def minify_json_for_llm(json_str):
    try:
        data = json.loads(json_str)
        return json.dumps(data, separators=(',', ':'))
    except:
        return json_str

def build_master_cv(resume_data):
    return f"""{resume_data.get('header', '')}

## SUMMARY
{{{{
{resume_data.get('summary', '')}
}}}}

## EDUCATION
{resume_data.get('education', '')}

## QUANTITATIVE & DATA PROJECTS
{{{{ }}}}

## PROFESSIONAL EXPERIENCE
{{{{
{resume_data.get('experience', '')}
}}}}

## TECHNICAL SKILLS
{resume_data.get('skills', '')}
"""

def build_soft_skills_text(skills_list):
    lines = ["# Soft Skills & Cognitive Abilities\n"]
    for item in skills_list:
        trait = item.get("trait", "").strip()
        example = item.get("example", "").strip()
        if trait or example:
            lines.append(f"* **{trait}:** {example}")
    return "\n\n".join(lines)

# --- 1. DEFINING THE AI AGENT'S STRUCTURED OUTPUT ---
class JobAnalysisAgentOutput(BaseModel):
    job_title: str = Field(description="Formal job title extracted.")
    company_name: str = Field(description="Company name extracted.")
    keywords: List[str] = Field(description="Key skills and requirements.")
    matches: List[str] = Field(description="Clear alignment points.")
    gaps: List[str] = Field(description="Missing skills or experience gaps.")
    relevant_repo_names: List[str] = Field(description="Exactly top 6 relevant project names. Return empty if skipping.")
    new_intro: str = Field(description="Optimized professional summary. NO Markdown headers.")
    cover_letter_body: str = Field(description="4-paragraph cover letter following strict prompt guidelines. NO hyphens.")
    success_probability: str = Field(description="Estimated match level (e.g., '85% Match').")

# --- 2. HYPER-OPTIMIZED 1-PAGE PDF RESUME ENGINE ---
class MarkdownResumePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_margins(12, 10, 12)
        self.add_page()
        
    def add_markdown_content(self, markdown_text):
        cleaned_text = sanitize_text(markdown_text)
        lines = cleaned_text.replace("{{", "").replace("}}", "").split("\n")
        
        for line in lines:
            line = line.strip()
            if not line:
                self.ln(0.6) 
                continue
                
            self.set_x(12)
                
            if line.startswith("# "):
                self.set_font("Helvetica", "B", 15)
                self.set_text_color(26, 54, 93) 
                self.cell(0, 6, sanitize_text(line[2:]), ln=True, align="C")
                
            elif line.startswith("## "):
                self.ln(1.2)
                self.set_font("Helvetica", "B", 11) 
                self.set_text_color(43, 108, 176) 
                self.cell(0, 4.5, sanitize_text(line[3:].upper()), ln=True)
                self.set_draw_color(226, 232, 240)
                self.line(12, self.get_y(), 198, self.get_y()) 
                self.ln(0.8)
                
            elif line.startswith("**"):
                self.ln(1.0) 
                clean_line = line.replace("**", "") 
                
                if " | " in clean_line:
                    parts = clean_line.split(" | ", 1)
                    title_part = sanitize_text(parts[0])
                    url_part = sanitize_text(parts[1])
                    
                    self.set_font("Helvetica", "B", 11) 
                    self.set_text_color(45, 55, 72)
                    
                    title_w = self.get_string_width(title_part + " | ")
                    self.cell(title_w, 4.2, title_part + " | ", ln=False)
                    
                    self.set_font("Helvetica", "I", 8)
                    self.set_text_color(100, 116, 139)
                    self.cell(0, 4.2, url_part, ln=True)
                else:
                    self.set_font("Helvetica", "B", 11) 
                    self.set_text_color(45, 55, 72) 
                    self.cell(0, 4.2, sanitize_text(clean_line), ln=True)
                
            elif line.startswith("*") or line.startswith("•") or line.startswith("-"):
                self.set_font("Helvetica", "", 8) 
                self.set_text_color(45, 55, 72)
                self.set_x(16) 
                bullet_content = sanitize_text(re.sub(r'^[*•-]\s*', '', line).replace("**", ""))
                
                if ":" in bullet_content:
                    parts = bullet_content.split(":", 1)
                    self.cell(3, 3.8, chr(149)) 
                    self.set_font("Helvetica", "B", 8)
                    width = self.get_string_width(parts[0] + ": ")
                    self.cell(width, 3.8, parts[0] + ": ")
                    self.set_font("Helvetica", "", 8)
                    self.multi_cell(0, 3.8, parts[1].strip())
                else:
                    self.cell(3, 3.8, chr(149)) 
                    self.multi_cell(0, 3.8, bullet_content) 
                
            else:
                self.set_text_color(45, 55, 72)
                safe_line = sanitize_text(line)
                if "@" in safe_line or "linkedin.com" in safe_line:
                    self.set_font("Helvetica", "", 8.5)
                    self.set_text_color(74, 85, 104) 
                    self.cell(0, 3.8, safe_line, ln=True, align="C")
                else:
                    self.set_font("Helvetica", "", 8)
                    self.multi_cell(0, 3.8, safe_line)

# --- 3. COMPACT DOCX RESUME ENGINE ---
def generate_docx(markdown_text, output_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    lines = sanitize_text(markdown_text).replace("{{", "").replace("}}", "").split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[2:].upper())
            run.bold = True
            run.font.size = Pt(14)
        elif "@" in line or "linkedin.com" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(74, 85, 104)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line[3:].upper())
            run.bold = True
            run.font.size = Pt(11) 
            run.font.color.rgb = RGBColor(43, 108, 176)
        elif line.startswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0.5)
            clean_line = line.replace("**", "")
            
            if " | " in clean_line:
                parts = clean_line.split(" | ", 1)
                r1 = p.add_run(parts[0] + " | ")
                r1.bold = True
                r1.font.size = Pt(10.5) 
                
                r2 = p.add_run(parts[1])
                r2.italic = True
                r2.font.size = Pt(8) 
                r2.font.color.rgb = RGBColor(100, 116, 139)
            else:
                r = p.add_run(clean_line)
                r.bold = True
                r.font.size = Pt(10.5) 
        elif line.startswith("* ") or line.startswith("• ") or line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0.5)
            clean_line = re.sub(r'^[*•-]\s*', '', line).replace("**", "")
            if ":" in clean_line:
                parts = clean_line.split(":", 1)
                r1 = p.add_run(parts[0] + ": ")
                r1.bold = True
                r1.font.size = Pt(8)
                r2 = p.add_run(parts[1].strip())
                r2.font.size = Pt(8) 
            else:
                r = p.add_run(clean_line)
                r.font.size = Pt(8) 
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(line)
            r.font.size = Pt(8)

    doc.save(output_path)

# --- 4. COVER LETTER PDF GENERATOR (PII SCRUBBED) ---
class StructuredCoverLetterPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_margins(20, 20, 20)
        self.add_page()
        
    def generate_layout(self, body_content, company_name):
        safe_body = sanitize_text(body_content)
        safe_company = sanitize_text(company_name) if company_name else ""
        
        self.set_font("Helvetica", "", 11)
        self.set_text_color(45, 55, 72)
        
        salutation = f"Dear Hiring Team at {safe_company}," if safe_company else "Dear Hiring Manager,"
        self.cell(0, 6, salutation, ln=True)
        self.ln(5)
        
        self.multi_cell(0, 6, safe_body)
        self.ln(6)
        
        self.cell(0, 6, "Thanks for your time,", ln=True)
        self.ln(4)
        
        self.set_font("Helvetica", "B", 11)
        self.cell(0, 5, "[YOUR NAME]", ln=True)
        self.set_font("Helvetica", "", 10)
        self.set_text_color(74, 85, 104)
        self.cell(0, 5, "[CITY, STATE] | [PHONE] | [EMAIL]", ln=True)
        self.ln(6)
        
        self.set_text_color(45, 55, 72)
        today_str = datetime.today().strftime('%B %d, %Y')
        self.cell(0, 5, today_str, ln=True)

# --- 5. COVER LETTER DOCX GENERATOR (PII SCRUBBED) ---
def generate_cover_letter_docx(body_content, company_name, output_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    safe_body = sanitize_text(body_content)
    safe_company = sanitize_text(company_name) if company_name else ""

    p_sal = doc.add_paragraph()
    salutation = f"Dear Hiring Team at {safe_company}," if safe_company else "Dear Hiring Manager,"
    p_sal.add_run(salutation)
    p_sal.paragraph_format.space_after = Pt(12)
    
    paragraphs = safe_body.split("\n\n")
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.15
            p.add_run(para.strip())
            
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_before = Pt(12)
    p_close.paragraph_format.space_after = Pt(18)
    p_close.add_run("Thanks for your time,")
    
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_after = Pt(4)
    r_name = p_sig.add_run("[YOUR NAME]")
    r_name.bold = True
    
    p_info = doc.add_paragraph()
    p_info.paragraph_format.space_after = Pt(12)
    p_info.add_run("[CITY, STATE] | [PHONE] | [EMAIL]")
    
    p_date = doc.add_paragraph()
    today_str = datetime.today().strftime('%B %d, %Y')
    p_date.add_run(today_str)
    
    doc.save(output_path)

# --- 6. HELPER FUNCTIONS ---
def extract_text_from_url(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        return soup.get_text(separator=' ', strip=True)
    except Exception as e:
        st.error(f"Scraper failed: {e}")
        return None

# CACHED API CALL 1: Core generation.
@st.cache_data(ttl=3600, show_spinner=False)
def call_ai_agent(api_key, jd_text, master_cv, github_json, soft_skills_text, custom_prompt, feedback=None, use_template=False):
    client = genai.Client(api_key=api_key)
    
    mini_jd = minify_text_for_llm(jd_text)
    mini_cv = minify_text_for_llm(master_cv)
    mini_soft = minify_text_for_llm(soft_skills_text)
    
    if use_template:
        mini_git = "[]"
        project_instruction = "2. SKIP PROJECT SORTING: The user is using a pre-saved template. Return an empty list for `relevant_repo_names`."
    else:
        mini_git = minify_json_for_llm(github_json)
        project_instruction = "2. SORT PROJECTS: Select EXACTLY top 6 relevant projects from the GitHub JSON."

    final_system_prompt = custom_prompt.replace("{project_instruction}", project_instruction)

    prompt = f"""
{final_system_prompt}

JD:
{mini_jd}

CV:
{mini_cv}

GITHUB:
{mini_git}

SOFT SKILLS:
{mini_soft}
    """
    
    if feedback:
        prompt += f"\n\nUSER FEEDBACK FOR REGENERATION:\n'{feedback}'\nAdjust analysis, projects, summary, and cover letter based on feedback."

    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=JobAnalysisAgentOutput,
                temperature=0.2 if not feedback else 0.4
            ),
        )
        return json.loads(response.text)
    except ClientError as e:
        if e.code == 429:
            st.error("⚠️ API Rate Limit Reached (429 Quota Exceeded). You have hit the free tier limit. Please wait a bit or upgrade your Google AI Studio billing plan.")
            st.stop()
        else:
            raise e

# CACHED API CALL 2: Cover Letter regeneration.
@st.cache_data(ttl=3600, show_spinner=False)
def call_cover_letter_regeneration(api_key, jd_text, master_cv, soft_skills_text, current_cl, feedback, custom_regen_prompt):
    client = genai.Client(api_key=api_key)
    
    mini_jd = minify_text_for_llm(jd_text)
    mini_cv = minify_text_for_llm(master_cv)
    mini_soft = minify_text_for_llm(soft_skills_text)

    prompt = f"""
    Regenerate cover letter based on feedback.
    JD: {mini_jd}
    CV: {mini_cv}
    SOFT SKILLS: {mini_soft}
    CURRENT CL: {current_cl}
    FEEDBACK: {feedback}

    {custom_regen_prompt}
    """
    
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.4
            ),
        )
        return response.text.strip()
    except ClientError as e:
        if e.code == 429:
            st.error("⚠️ API Rate Limit Reached (429 Quota Exceeded). Please wait.")
            st.stop()
        else:
            raise e

# --- 7. STREAMLIT UI & TAB LOGIC ---
st.set_page_config(page_title="AI Career Agent", layout="wide")

# --- AUTHENTICATION FALLBACK LOGIC ---
env_api_key = os.getenv("GEMINI_API_KEY")
secrets_api_key = None
try:
    if "GEMINI_API_KEY" in st.secrets:
        secrets_api_key = st.secrets["GEMINI_API_KEY"]
except Exception:
    pass

default_key = env_api_key or secrets_api_key or ""

st.sidebar.header("🔑 Authentication")
user_api_key = st.sidebar.text_input(
    "Google Gemini API Key", 
    value=default_key, 
    type="password", 
    help="Loaded automatically from .env or Secrets if available."
)

if not user_api_key:
    st.sidebar.warning("Please enter your API Key to enable the Agent.")

st.title("💼 Tailored AI Career Agent")

# Define Tabs
tab1, tab2, tab3, tab4, tab5 = st.tabs(["📄 Job Description & Agent", "📝 CV", "📂 Projects", "💡 Soft Skills", "⚙️ AI Prompts"])

# --- TAB 2: CV ---
with tab2:
    st.subheader("Edit Master Resume Sections")
    st.markdown("Edit the core sections of your resume below. The **Projects** section is dynamically injected from the 'Projects' tab during generation.")
    
    if "resume_data" not in st.session_state:
        with open("profile_data/resume_data.json", "r") as f:
            st.session_state.resume_data = json.load(f)
            
    r_data = st.session_state.resume_data
    r_data['header'] = st.text_area("Contact & Header", r_data.get('header', ''), height=100)
    r_data['summary'] = st.text_area("Professional Summary", r_data.get('summary', ''), height=150)
    r_data['education'] = st.text_area("Education", r_data.get('education', ''), height=150)
    r_data['experience'] = st.text_area("Professional Experience", r_data.get('experience', ''), height=250)
    r_data['skills'] = st.text_area("Technical Skills", r_data.get('skills', ''), height=200)

    if st.button("💾 Save Resume Sections"):
        with open("profile_data/resume_data.json", "w") as f:
            json.dump(st.session_state.resume_data, f, indent=4)
        st.cache_data.clear() 
        st.success("Master CV Sections Updated!")

# --- TAB 3: PROJECTS ---
with tab3:
    st.subheader("Edit GitHub Projects")
    
    if "projects_data" not in st.session_state:
        with open("profile_data/github_repos.json", "r") as f:
            st.session_state.projects_data = json.load(f)

    def add_project():
        st.session_state.projects_data.append({
            "name": "New Project", 
            "url": "", 
            "description": "", 
            "tech_stack": [], 
            "bullets": []
        })

    def delete_project(idx):
        st.session_state.projects_data.pop(idx)

    st.button("➕ Add New Project", on_click=add_project)
    
    for i, proj in enumerate(st.session_state.projects_data):
        with st.expander(f"{proj.get('name', 'Untitled Project')}", expanded=False):
            proj['name'] = st.text_input("Name", proj.get('name', ''), key=f"p_name_{i}")
            proj['url'] = st.text_input("URL", proj.get('url', ''), key=f"p_url_{i}")
            proj['description'] = st.text_area("Description", proj.get('description', ''), key=f"p_desc_{i}")
            
            tech_str = st.text_input("Tech Stack (comma-separated)", ", ".join(proj.get('tech_stack', [])), key=f"p_tech_{i}")
            proj['tech_stack'] = [t.strip() for t in tech_str.split(",") if t.strip()]
            
            bull_str = st.text_area("Bullets (one per line)", "\n".join(proj.get('bullets', [])), key=f"p_bull_{i}")
            proj['bullets'] = [b.strip() for b in bull_str.split("\n") if b.strip()]
            
            st.button("❌ Delete Project", key=f"p_del_{i}", on_click=delete_project, args=(i,))

    st.divider()
    if st.button("💾 Save All Projects"):
        with open("profile_data/github_repos.json", "w") as f:
            json.dump(st.session_state.projects_data, f, indent=4)
        st.cache_data.clear() 
        st.success("Projects JSON Updated Successfully!")

# --- TAB 4: SOFT SKILLS (USER FRIENDLY EXPANDERS) ---
with tab4:
    st.subheader("Edit Soft Skills & Cognitive Abilities")
    st.markdown("Add or edit your soft skills, cognitive traits, and real-world examples below. These will be dynamically woven into Paragraph 2 of your cover letter.")
    
    if "soft_skills_data" not in st.session_state:
        if os.path.exists("profile_data/soft_skills.json"):
            with open("profile_data/soft_skills.json", "r") as f:
                st.session_state.soft_skills_data = json.load(f)
        else:
            st.session_state.soft_skills_data = []

    def add_soft_skill():
        st.session_state.soft_skills_data.append({
            "trait": "New Cognitive Trait", 
            "example": ""
        })

    def delete_soft_skill(idx):
        st.session_state.soft_skills_data.pop(idx)

    st.button("➕ Add New Soft Skill", on_click=add_soft_skill)
    
    for i, skill in enumerate(st.session_state.soft_skills_data):
        with st.expander(f"{skill.get('trait', 'Untitled Skill')}", expanded=False):
            skill['trait'] = st.text_input("Trait / Skill Title", skill.get('trait', ''), key=f"s_trait_{i}")
            skill['example'] = st.text_area("Example / Real-world Experience", skill.get('example', ''), key=f"s_ex_{i}")
            st.button("❌ Delete Soft Skill", key=f"s_del_{i}", on_click=delete_soft_skill, args=(i,))

    st.divider()
    if st.button("💾 Save All Soft Skills"):
        with open("profile_data/soft_skills.json", "w") as f:
            json.dump(st.session_state.soft_skills_data, f, indent=4)
            
        # Also sync to markdown format
        with open("profile_data/soft_skills.md", "w") as f:
            f.write(build_soft_skills_text(st.session_state.soft_skills_data))
            
        st.cache_data.clear() 
        st.success("Soft Skills Updated Successfully!")

# --- TAB 5: AI PROMPTS ---
with tab5:
    st.subheader("⚙️ Prompt Engineering")
    st.write("Adjust the instructions sent to the AI Agent. Note: Keep the `{project_instruction}` tag in the Main Prompt.")
    
    if "prompts_data" not in st.session_state:
        with open("profile_data/prompts.json", "r") as f:
            st.session_state.prompts_data = json.load(f)
            
    p_data = st.session_state.prompts_data
    p_data['main_prompt'] = st.text_area("Main Generation Prompt (CV & Cover Letter)", p_data.get('main_prompt', ''), height=450)
    p_data['regen_prompt'] = st.text_area("Cover Letter Regeneration Prompt", p_data.get('regen_prompt', ''), height=300)

    if st.button("💾 Save Prompts"):
        with open("profile_data/prompts.json", "w") as f:
            json.dump(st.session_state.prompts_data, f, indent=4)
        st.cache_data.clear() 
        st.success("AI Prompts Updated Successfully!")

# --- TAB 1: JOB DESCRIPTION & AGENT ---
with tab1:
    if "agent_results" not in st.session_state:
        st.session_state.agent_results = None
    if "current_jd" not in st.session_state:
        st.session_state.current_jd = ""
    if "is_approved" not in st.session_state:
        st.session_state.is_approved = False
    if "show_regeneration" not in st.session_state:
        st.session_state.show_regeneration = False
    if "show_cl_regeneration" not in st.session_state:
        st.session_state.show_cl_regeneration = False

    use_template = False
    template_content = ""
    template_files = os.listdir("templates") if os.path.exists("templates") else []
    if template_files:
        selected_template = st.selectbox("Or load a previously saved Template:", ["-- Select --"] + template_files)
        if selected_template != "-- Select --":
            use_template = True
            with open(os.path.join("templates", selected_template), "r") as f:
                template_content = f.read()
            st.info(f"Loaded Template: {selected_template}. Custom project selection from JSON will be bypassed.")

    input_type = st.radio("Input Method:", ["Paste Job Description", "Provide URL"])
    job_description = ""
    
    if input_type == "Paste Job Description":
        job_description = st.text_area("Paste JD Here:", height=200)
    else:
        job_url = st.text_input("Enter URL:")
        if job_url:
            job_description = extract_text_from_url(job_url)

    if st.button("🚀 Analyze Job & Generate Profile") and job_description:
        if not user_api_key:
            st.error("Please enter your Google Gemini API Key in the sidebar to proceed.")
        else:
            with st.spinner("Analyzing Job & Optimizing Profile (Cached to save API)..."):
                st.session_state.current_jd = job_description
                st.session_state.use_template = use_template
                st.session_state.template_content = template_content
                
                with open("profile_data/resume_data.json", "r") as f:
                    r_data = json.load(f)
                master_cv_content = build_master_cv(r_data)
                
                with open("profile_data/github_repos.json", "r") as f:
                    github_repos_content = f.read()
                    
                if os.path.exists("profile_data/soft_skills.json"):
                    with open("profile_data/soft_skills.json", "r") as f:
                        s_data = json.load(f)
                    soft_skills_text = build_soft_skills_text(s_data)
                else:
                    with open("profile_data/soft_skills.md", "r") as f:
                        soft_skills_text = f.read()
                    
                with open("profile_data/prompts.json", "r") as f:
                    loaded_prompts = json.load(f)
                
                st.session_state.agent_results = call_ai_agent(
                    api_key=user_api_key,
                    jd_text=job_description, 
                    master_cv=template_content if use_template else master_cv_content, 
                    github_json=github_repos_content, 
                    soft_skills_text=soft_skills_text,
                    custom_prompt=loaded_prompts.get('main_prompt', ''),
                    use_template=use_template
                )
                
                st.session_state.is_approved = False
                st.session_state.show_regeneration = False
                st.session_state.show_cl_regeneration = False
                st.rerun()

    if st.session_state.agent_results:
        res = st.session_state.agent_results
        active_use_template = st.session_state.get('use_template', False)
        
        with open("profile_data/resume_data.json", "r") as f:
            cv_base = st.session_state.get('template_content', "") if active_use_template else build_master_cv(json.load(f))
        
        st.subheader(f"Analysis for: {res['job_title']} at {res['company_name']}")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Success Probability", res['success_probability'])
            st.write("**Keywords:**", ", ".join(res['keywords']))
        with col2:
            st.write("✅ **Matches:**")
            for m in res['matches']: st.write(f"- {m}")
        with col3:
            st.write("❌ **Gaps:**")
            for g in res['gaps']: st.write(f"- {g}")
            
        st.divider()
        
        clean_intro = re.sub(r'^#+.*?\n', '', res['new_intro'] + '\n', flags=re.IGNORECASE).strip()
        
        if active_use_template:
            preview_cv = re.sub(r'(##\s*SUMMARY\s*\n).*?(?=\n##\s)', f"\\1{clean_intro}\n\n", cv_base, flags=re.DOTALL | re.IGNORECASE)
        else:
            with open("profile_data/github_repos.json", "r") as f:
                github_data = json.load(f)
            
            project_markdown_blocks = []
            for repo_name in res['relevant_repo_names'][:6]:  
                match = next((r for r in github_data if r['name'].lower() == repo_name.lower()), None)
                if match:
                    bullets_str = "\n".join([f"* {b}" for b in match['bullets']])
                    primary_tech = match.get('tech_stack', [''])[0] if match.get('tech_stack') else ""
                    tech_str = f" ({primary_tech})" if primary_tech else ""
                    repo_url = match.get('url', '[YOUR GITHUB URL]')
                    block = f"**{match['name']}{tech_str} | {repo_url}**\n{bullets_str}\n"
                    project_markdown_blocks.append(block)
            
            final_projects_section = "\n".join(project_markdown_blocks)
            preview_cv = cv_base
            preview_cv = re.sub(r'##\s*SUMMARY\s*\{\{.*?\}\}', f"## SUMMARY\n{clean_intro}", preview_cv, flags=re.DOTALL | re.IGNORECASE)
            preview_cv = re.sub(r'##\s*QUANTITATIVE & DATA PROJECTS\s*\{\{.*?\}\}', f"## QUANTITATIVE & DATA PROJECTS\n{final_projects_section}", preview_cv, flags=re.DOTALL | re.IGNORECASE)
            preview_cv = preview_cv.replace("{{", "").replace("}}", "")
        
        with st.expander("👀 Preview Generated CV", expanded=True):
            st.markdown(preview_cv)
            
        st.divider()
        
        st.write("### Review & Approve")
        st.checkbox("Review Match & Gap Analysis", value=st.session_state.is_approved)
        st.checkbox("Review Sorted Projects", value=st.session_state.is_approved)
        st.checkbox("Review Optimized Summary", value=st.session_state.is_approved)
        
        col_btn1, col_btn2 = st.columns(2)
        
        with col_btn1:
            if st.button("✅ Approve All & Continue"):
                st.session_state.is_approved = True
                st.session_state.show_regeneration = False
                st.rerun()
                    
        with col_btn2:
            if st.button("🔄 Regenerate"):
                st.session_state.show_regeneration = True
                st.session_state.is_approved = False
                st.rerun()
                
        if st.session_state.show_regeneration:
            st.warning("You selected Regenerate.")
            regen_feedback = st.text_area("What is the problem? (e.g., 'Make the summary shorter', 'Swap project X with Y')")
            if st.button("Submit Feedback & Regenerate"):
                if not user_api_key:
                    st.error("Please enter your API Key in the sidebar.")
                else:
                    with st.spinner("Regenerating Profile based on feedback..."):
                        if os.path.exists("profile_data/soft_skills.json"):
                            with open("profile_data/soft_skills.json", "r") as f:
                                s_data = json.load(f)
                            current_soft_skills = build_soft_skills_text(s_data)
                        else:
                            with open("profile_data/soft_skills.md", "r") as f:
                                current_soft_skills = f.read()
                                
                        with open("profile_data/github_repos.json", "r") as f:
                            github_repos_content = f.read()
                        with open("profile_data/prompts.json", "r") as f:
                            loaded_prompts = json.load(f)
                            
                        st.session_state.agent_results = call_ai_agent(
                            api_key=user_api_key,
                            jd_text=st.session_state.current_jd, 
                            master_cv=cv_base, 
                            github_json=github_repos_content, 
                            soft_skills_text=current_soft_skills,
                            custom_prompt=loaded_prompts.get('main_prompt', ''),
                            feedback=regen_feedback,
                            use_template=active_use_template
                        )
                        st.session_state.show_regeneration = False
                        st.session_state.is_approved = False
                        st.rerun()

        if st.session_state.is_approved:
            st.success("All sections approved! Choose an export format below:")
            
            safe_company = res['company_name'].replace(' ', '_').replace('/', '_')
            safe_title = res['job_title'].replace(' ', '_').replace('/', '_')
            
            col_a, col_b, col_c, col_d = st.columns(4)
            
            with col_a:
                if st.button("A. CV (PDF)"):
                    pdf_path = f"output/{safe_company}_{safe_title}_Resume.pdf"
                    pdf_gen = MarkdownResumePDF()
                    pdf_gen.add_markdown_content(preview_cv)
                    pdf_gen.output(pdf_path)
                    with open(pdf_path, "rb") as f:
                        st.download_button("Download PDF", f, file_name=f"{safe_company}_CV.pdf", mime="application/pdf")

            with col_b:
                if st.button("B. CV (DOCX)"):
                    docx_path = f"output/{safe_company}_{safe_title}_Resume.docx"
                    generate_docx(preview_cv, docx_path)
                    with open(docx_path, "rb") as f:
                        st.download_button("Download DOCX", f, file_name=f"{safe_company}_CV.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            with col_c:
                if st.button("C. Save Template"):
                    template_path = f"templates/{safe_company}_Template.md"
                    with open(template_path, "w") as f:
                        f.write(preview_cv)
                    st.success(f"Saved to {template_path}!")

            st.divider()
            st.subheader("✉️ Structured Cover Letter Generator")
            st.write("Previewing the generated 4-paragraph structured cover letter:")
            
            st.info(res['cover_letter_body'])
            
            cl_col1, cl_col2 = st.columns(2)
            with cl_col1:
                if st.button("🔄 Regenerate Cover Letter"):
                    st.session_state.show_cl_regeneration = True
                    st.rerun()
            
            if st.session_state.get("show_cl_regeneration", False):
                cl_feedback = st.text_area("What is the problem with the cover letter? (e.g., 'Make paragraph 2 more direct', 'Change tone')")
                if st.button("Submit Cover Letter Feedback & Regenerate"):
                    if not user_api_key:
                        st.error("Please enter your API Key in the sidebar.")
                    else:
                        with st.spinner("Regenerating Cover Letter..."):
                            if os.path.exists("profile_data/soft_skills.json"):
                                with open("profile_data/soft_skills.json", "r") as f:
                                    s_data = json.load(f)
                                current_soft_skills = build_soft_skills_text(s_data)
                            else:
                                with open("profile_data/soft_skills.md", "r") as f:
                                    current_soft_skills = f.read()
                                    
                            with open("profile_data/prompts.json", "r") as f:
                                loaded_prompts = json.load(f)
                                
                            new_cl = call_cover_letter_regeneration(
                                api_key=user_api_key,
                                jd_text=st.session_state.current_jd,
                                master_cv=cv_base,
                                soft_skills_text=current_soft_skills,
                                current_cl=res['cover_letter_body'],
                                feedback=cl_feedback,
                                custom_regen_prompt=loaded_prompts.get('regen_prompt', '')
                            )
                            st.session_state.agent_results['cover_letter_body'] = new_cl
                            st.session_state.show_cl_regeneration = False
                            st.rerun()

            cl_dl_col1, cl_dl_col2 = st.columns(2)
            with cl_dl_col1:
                if st.button("Download Cover Letter (PDF)"):
                    cl_path = f"output/{safe_company}_{safe_title}_Cover_Letter.pdf"
                    cl_gen = StructuredCoverLetterPDF()
                    cl_gen.generate_layout(res['cover_letter_body'], res['company_name'])
                    cl_gen.output(cl_path)
                    with open(cl_path, "rb") as f:
                        st.download_button("Download CL PDF", f, file_name=f"{safe_company}_Cover_Letter.pdf", mime="application/pdf")
            with cl_dl_col2:
                if st.button("Download Cover Letter (DOCX)"):
                    cl_docx_path = f"output/{safe_company}_{safe_title}_Cover_Letter.docx"
                    generate_cover_letter_docx(res['cover_letter_body'], res['company_name'], cl_docx_path)
                    with open(cl_docx_path, "rb") as f:
                        st.download_button("Download CL DOCX", f, file_name=f"{safe_company}_Cover_Letter.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
